"""Renders a generated cover letter into the standard Word template: contact header, date,
recipient block, "Re:" line, body, sign-off. Any field the app has no real data for (hiring
manager's name, company address, ...) is written as an <insert xx> placeholder for the user to
fill in by hand, rather than guessed or left blank.
"""

import re
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

BLOCK_GAP = Pt(9)
LINE_GAP = Pt(0)


def _add_paragraph(doc: Document, text: str, bold: bool = False, space_after: Pt = LINE_GAP):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = space_after
    fmt.line_spacing = 1.08
    return paragraph


def _safe_filename_part(text: str) -> str:
    # Strips characters Windows/macOS forbid in filenames; collapses whatever whitespace is left.
    cleaned = re.sub(r'[\\/:*?"<>|]+', "", text)
    return re.sub(r"\s+", " ", cleaned).strip()


def build_export_filename(name: str, position: str | None) -> str:
    name_part = _safe_filename_part(name)
    position_part = _safe_filename_part(position or "")
    base = f"{name_part} {position_part}".strip() if name_part else position_part
    return f"{base + ' ' if base else ''}Cover Letter.docx"


def build_cover_letter_docx(
    basic_info: dict, company: str | None, position: str | None, letter_text: str
) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    name = (basic_info.get("full_name") or "").strip() or "<Your Name>"
    phone = (basic_info.get("phone") or "").strip() or "<Your Contact Number>"
    email = (basic_info.get("email") or "").strip() or "<Your Email Address>"
    address = (basic_info.get("location") or "").strip() or "<Your Home Address>"
    company_name = (company or "").strip() or "<Company Name>"
    role_title = (position or "").strip() or "<Role Title>"

    _add_paragraph(doc, name)
    _add_paragraph(doc, phone)
    _add_paragraph(doc, email)
    _add_paragraph(doc, address, space_after=BLOCK_GAP)

    _add_paragraph(doc, datetime.now().strftime("%d %B %Y"), space_after=BLOCK_GAP)

    _add_paragraph(doc, "<Hiring Manager's Name>")
    _add_paragraph(doc, company_name)
    _add_paragraph(doc, "<Company Address>", space_after=BLOCK_GAP)

    _add_paragraph(doc, "Dear <Hiring Manager's Name>,", bold=True, space_after=BLOCK_GAP)

    _add_paragraph(doc, f"Re: {role_title} position", bold=True, space_after=BLOCK_GAP)

    paragraphs = [p.strip() for p in letter_text.strip().split("\n\n") if p.strip()]
    for para in paragraphs:
        _add_paragraph(doc, para, space_after=BLOCK_GAP)

    _add_paragraph(doc, "Yours sincerely,", bold=True, space_after=BLOCK_GAP)
    _add_paragraph(doc, name)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
